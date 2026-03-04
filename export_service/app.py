# app.py
# =============================================================
# Engine v3 – Vision Pipeline  +  Universal DOCX Renderer
# Zero hallucination: AI outputs JSON, engine renders it
#
# .env:
#   OPENROUTER_API_KEY=...
#   HANDW_API_KEY=...
#   GOOGLE_VISION_API_KEY=...
#   REDIS_URL=redis://...
#   OCR_MODEL=google/gemini-2.0-flash-001
# =============================================================

import os
import base64
import json
import re
import io
import unicodedata
import requests
import traceback
import fitz
import time
import numpy as np
import cv2
import redis as redis_lib

from fastapi import FastAPI, UploadFile, File, HTTPException, Form, BackgroundTasks, Request
from fastapi.responses import StreamingResponse, JSONResponse, Response
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel, Field
from typing import Any, Optional
from datetime import datetime

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from dotenv import load_dotenv
load_dotenv()


# ─────────────────────────────────────────────────────────────
# GLOBAL CONFIG
# ─────────────────────────────────────────────────────────────

ENGINE_VERSION     = "v3.0.0"
OPENROUTER_API_KEY = os.getenv("OPENROUTER_API_KEY")
OPENROUTER_URL     = "https://openrouter.ai/api/v1/chat/completions"
MODEL              = os.getenv("OCR_MODEL", "google/gemini-2.0-flash-001")
MAX_PDF_PAGES      = 20

OCR_HEADERS = {
    "Authorization": f"Bearer {OPENROUTER_API_KEY}",
    "Content-Type":  "application/json",
    "HTTP-Referer":  "http://localhost",
    "X-Title":       "Doc-Reconstructor-v3",
}

if not OPENROUTER_API_KEY:
    raise RuntimeError("OPENROUTER_API_KEY not set")

BASE_DIR = os.path.dirname(__file__)

API_KEY = os.getenv("HANDW_API_KEY")
if not API_KEY:
    raise RuntimeError("HANDW_API_KEY is NOT set.")


# ─────────────────────────────────────────────────────────────
# FASTAPI APP + MIDDLEWARE
# ─────────────────────────────────────────────────────────────

app = FastAPI(title="Handwritten-to-Doc Engine v3")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

@app.middleware("http")
async def api_key_guard(request: Request, call_next):
    if request.url.path in ["/docs", "/openapi.json", "/redoc"]:
        return await call_next(request)
    if request.url.path.startswith("/api") or request.url.path == "/generate-docx":
        key = request.headers.get("x-api-key")
        if not API_KEY or key != API_KEY:
            return JSONResponse(status_code=401, content={"error": "UNAUTHORIZED"})
    return await call_next(request)


# ─────────────────────────────────────────────────────────────
# REDIS JOB STORE
# ─────────────────────────────────────────────────────────────

JOB_TTL = 60 * 60 * 3

def _get_redis():
    url = os.getenv("REDIS_URL")
    if not url:
        raise RuntimeError("REDIS_URL env var not set")
    return redis_lib.from_url(url, decode_responses=True)

def load_job(jobId: str):
    try:
        r   = _get_redis()
        raw = r.get(f"job:{jobId}")
        return json.loads(raw) if raw else None
    except Exception as e:
        log("⚠️ Redis load_job error", repr(e))
        return None

def update_job(jobId: str, **updates):
    try:
        r        = _get_redis()
        key      = f"job:{jobId}"
        raw      = r.get(key)
        existing = json.loads(raw) if raw else {"jobId": jobId}
        existing.update(updates)
        r.setex(key, JOB_TTL, json.dumps(existing))
    except Exception as e:
        log("⚠️ Redis update_job error", repr(e))


# ─────────────────────────────────────────────────────────────
# LOGGING
# ─────────────────────────────────────────────────────────────

def log(step, data=None):
    ts = datetime.now().strftime("%H:%M:%S")
    print(f"\n🟦 [{ts}] STEP: {step}")
    if data is not None:
        print(data)


# =============================================================
# ░░░░  SECTION 1 — OCR / VISION PIPELINE  ░░░░░░░░░░░░░░░░░░
# =============================================================

def is_pdf(data: bytes) -> bool:
    return data[:4] == b"%PDF"

def pdf_page_to_image_bytes(page) -> bytes:
    pix = page.get_pixmap(dpi=300)
    img = np.frombuffer(pix.samples, dtype=np.uint8).reshape(pix.height, pix.width, pix.n)
    if pix.n == 4:
        img = cv2.cvtColor(img, cv2.COLOR_BGRA2BGR)
    ok, buf = cv2.imencode(".png", img)
    if not ok:
        raise ValueError("Failed to encode PDF page as PNG")
    return buf.tobytes()

def pdf_to_image_bytes(pdf_bytes: bytes) -> bytes:
    doc         = fitz.open(stream=pdf_bytes, filetype="pdf")
    total_pages = min(len(doc), MAX_PDF_PAGES)
    log("PDF pages", f"{total_pages} / {len(doc)}")
    page_images = []
    for i in range(total_pages):
        raw = pdf_page_to_image_bytes(doc.load_page(i))
        page_images.append(cv2.imdecode(np.frombuffer(raw, np.uint8), cv2.IMREAD_COLOR))
    if len(page_images) == 1:
        stitched = page_images[0]
    else:
        max_w  = max(img.shape[1] for img in page_images)
        padded = []
        for img in page_images:
            h, w = img.shape[:2]
            if w < max_w:
                img = np.hstack([img, np.ones((h, max_w - w, 3), dtype=np.uint8) * 255])
            padded.append(img)
        sep         = np.ones((10, max_w, 3), dtype=np.uint8) * 255
        interleaved = []
        for i, img in enumerate(padded):
            interleaved.append(img)
            if i < len(padded) - 1:
                interleaved.append(sep)
        stitched = np.vstack(interleaved)
    ok, buf = cv2.imencode(".png", stitched)
    if not ok:
        raise ValueError("Failed to encode stitched PDF as PNG")
    return buf.tobytes()

def to_png_bytes(raw_bytes: bytes) -> bytes:
    img = cv2.imdecode(np.frombuffer(raw_bytes, np.uint8), cv2.IMREAD_COLOR)
    if img is None:
        raise ValueError("Cannot decode image")
    ok, buf = cv2.imencode(".png", img)
    if not ok:
        raise ValueError("Failed to re-encode image as PNG")
    return buf.tobytes()


# ─────────────────────────────────────────────────────────────
# STAGE 1 — Google Vision OCR  →  raw text only
# ─────────────────────────────────────────────────────────────

def stage1_extract_text(image_bytes: bytes) -> str:
    log("STAGE 1 — Google Vision OCR")
    t0      = time.time()
    api_key = os.getenv("GOOGLE_VISION_API_KEY")
    if not api_key:
        raise RuntimeError("GOOGLE_VISION_API_KEY not set")

    b64     = base64.b64encode(image_bytes).decode("utf-8")
    payload = {"requests": [{"image": {"content": b64},
                              "features": [{"type": "DOCUMENT_TEXT_DETECTION"}]}]}
    res = requests.post(
        f"https://vision.googleapis.com/v1/images:annotate?key={api_key}",
        json=payload, timeout=60)
    res.raise_for_status()

    raw_text = res.json()["responses"][0].get("fullTextAnnotation", {}).get("text", "")
    if not raw_text.strip():
        raise ValueError("Google Vision returned empty text")

    log("STAGE 1 done", f"{round(time.time()-t0,2)}s | {len(raw_text)} chars")
    return raw_text


# ─────────────────────────────────────────────────────────────
# STAGE 2 — Auditor  →  corrected raw text
# ─────────────────────────────────────────────────────────────

def extract_json_safe(text: str) -> dict:
    text = re.sub(r"^```(?:json)?\s*", "", text.strip())
    text = re.sub(r"\s*```$", "", text)
    try:
        return json.loads(text)
    except json.JSONDecodeError:
        log("⚠️ JSON parse failed")
        return {}

def stage2_audit(raw_text: str) -> dict:
    log("STAGE 2 — Auditor")
    t0 = time.time()

    prompt = f"""You are a strict document auditor. Analyze the OCR-extracted text below and return a JSON report.

TEXT:
```
{raw_text}
```

Check for:
1. Impossible dates
2. Broken/truncated words or sentences
3. Duplicate paragraphs
4. Numbers that look wrong
5. [?] placeholders — note locations
6. Text that looks invented rather than transcribed
7. Transposed digits in amounts
8. Sentences that end unnaturally — flag as POTENTIAL_FABRICATION

FABRICATION REMOVAL RULE:
If the LAST sentence ends abruptly or appears invented:
- Mark as POTENTIAL_FABRICATION
- In corrected_text: REMOVE it entirely, end with [DOCUMENT TRUNCATED]
- It is better to truncate than to present invented content

Return ONLY this JSON (no code fences):
{{
  "hallucination_risk": "low" | "medium" | "high",
  "issues_found": [],
  "corrections_made": [],
  "corrected_text": "<full corrected text — remove POTENTIAL_FABRICATION>"
}}"""

    payload = {
        "model": MODEL, "temperature": 0, "max_tokens": 4000,
        "messages": [
            {"role": "system", "content": "Return clean JSON only."},
            {"role": "user",   "content": prompt},
        ],
    }
    res    = requests.post(OPENROUTER_URL, headers=OCR_HEADERS, json=payload, timeout=90)
    res.raise_for_status()
    result = extract_json_safe(res.json()["choices"][0]["message"]["content"])
    risk   = result.get("hallucination_risk", "?").upper()
    log(f"STAGE 2 {'🟢' if risk=='LOW' else '🟡' if risk=='MEDIUM' else '🔴'}",
        f"{round(time.time()-t0,2)}s | risk={risk}")
    return result


# ─────────────────────────────────────────────────────────────
# STAGE 3 — Convert verified text → Universal Engine JSON
# ─────────────────────────────────────────────────────────────

def stage3_to_engine_json(verified_text: str) -> dict:
    log("STAGE 3 — Universal Engine JSON")
    t0 = time.time()

    prompt = f"""Convert the document text below into this exact JSON schema for rendering.

OUTPUT SCHEMA:
{{
  "title": "DOCUMENT TITLE (null if none)",
  "title_align": "center",
  "font": "Times New Roman",
  "font_size": 12,
  "page_size": "letter",
  "page_border": null,
  "header_text": null,
  "footer_text": null,
  "content": [ ...blocks ]
}}

BLOCK TYPES — use exactly these:

HEADING:
  {{ "type": "heading", "text": "SECTION NAME", "level": 1 }}
  level 1 = bold+underline, level 2 = bold, level 3 = bold+italic

PARAGRAPH:
  {{ "type": "paragraph", "align": "justify", "runs": [ ...runs ] }}

BULLET (dash list):
  {{ "type": "bullet", "runs": [ ...runs ] }}

NUMBERED LIST:
  {{ "type": "numbered", "index": 1, "runs": [ ...runs ] }}

TABLE:
  {{ "type": "table", "header": true, "rows": [["Col1","Col2"],["val","val"]] }}

SIGNATURE BLOCK (ONLY if signatures explicitly appear in document):
  {{ "type": "signature", "labels": ["Party A", "Party B"] }}

PAGE BREAK:
  {{ "type": "pagebreak" }}

HORIZONTAL RULE:
  {{ "type": "hr" }}

SPACER:
  {{ "type": "spacer", "size": 12 }}

RUN TYPES (inline, inside "runs" arrays):
  Plain:      {{ "text": "hello" }}
  Bold:       {{ "text": "hello", "bold": true }}
  Italic:     {{ "text": "hello", "italic": true }}
  Underline:  {{ "text": "hello", "underline": true }}
  Color:      {{ "text": "hello", "color": "#FF0000" }}
  Size:       {{ "text": "hello", "size": 14 }}
  Blank field:{{ "blank": true }}   ← use for ________ fill-in lines
  Line break: {{ "break": true }}

CRITICAL RULES — NO HALLUCINATION:
- Copy text EXACTLY as given. Do NOT add, invent, or complete anything.
- Blank lines (______) in the source = {{ "blank": true }} run in JSON.
- Square brackets like [Your Name] = copy exactly as plain text run.
- ONLY add "signature" block if signatures actually appear in the document text.
- Do NOT add content that is not in the source text.
- Output ONLY the JSON object. No explanation. No code fences.

DOCUMENT TEXT:
{verified_text}"""

    payload = {
        "model": MODEL, "temperature": 0, "max_tokens": 6000,
        "messages": [
            {"role": "system", "content": "You are a document structure extractor. Output only valid JSON matching the schema exactly."},
            {"role": "user",   "content": prompt},
        ],
    }
    res = requests.post(OPENROUTER_URL, headers=OCR_HEADERS, json=payload, timeout=90)
    res.raise_for_status()
    doc = extract_json_safe(res.json()["choices"][0]["message"]["content"])
    log("STAGE 3 done", f"{round(time.time()-t0,2)}s")
    return doc


# ─────────────────────────────────────────────────────────────
# PIPELINE ORCHESTRATOR
# ─────────────────────────────────────────────────────────────

def strip_truncated(text: str) -> str:
    marker = "[DOCUMENT TRUNCATED]"
    if marker in text:
        text = text[:text.index(marker)].rstrip()
        log("⚠️ TRUNCATION MARKER found — content cut")
    return text

def parse_document(image_bytes: bytes) -> dict:
    try:
        log("START parse_document", f"bytes={len(image_bytes)}")
        t0 = time.time()

        # Stage 1: Pure OCR — no AI interpretation
        raw_text = stage1_extract_text(image_bytes)
        raw_text = strip_truncated(raw_text)

        # Stage 2: Audit — fix errors, remove fabrications
        audit          = stage2_audit(raw_text)
        risk           = audit.get("hallucination_risk", "low")
        verified_text  = audit.get("corrected_text") or raw_text
        if not verified_text.strip():
            verified_text = raw_text
        verified_text  = strip_truncated(verified_text)

        # Stage 3: Structure — text → engine JSON
        engine_json    = stage3_to_engine_json(verified_text)

        total = round(time.time() - t0, 2)
        log("SUCCESS parse_document", f"total={total}s | risk={risk}")

        # Attach audit metadata
        engine_json["_audit"] = {
            "hallucination_risk": risk,
            "issues_found":       audit.get("issues_found", []),
            "corrections_made":   audit.get("corrections_made", []),
            "pipeline_seconds":   total,
            "engine_version":     ENGINE_VERSION,
        }
        return engine_json

    except Exception as e:
        log("❌ ERROR in parse_document", repr(e))
        traceback.print_exc()
        raise

def run_ocr_job(jobId: str):
    try:
        log("JOB START", jobId)
        job = load_job(jobId)
        if not job:
            raise RuntimeError("Job not found in Redis")
        update_job(jobId, state="processing")

        file_path = job.get("filePath")
        if not file_path or not os.path.exists(file_path):
            raise RuntimeError(f"File not found: {file_path!r}")

        with open(file_path, "rb") as f:
            raw_bytes = f.read()

        image_bytes = pdf_to_image_bytes(raw_bytes) if is_pdf(raw_bytes) else to_png_bytes(raw_bytes)
        document    = parse_document(image_bytes)

        update_job(jobId, state="ready", contentJson=document)
        log("JOB DONE", jobId)
    except Exception as e:
        log("JOB ERROR", repr(e))
        traceback.print_exc()
        update_job(jobId, state="error", detail=repr(e))


# =============================================================
# ░░░░  SECTION 2 — UNIVERSAL DOCX RENDER ENGINE  ░░░░░░░░░░░
# =============================================================

# ── Constants ─────────────────────────────────────────────────

ALIGN_MAP = {
    "left":    WD_ALIGN_PARAGRAPH.LEFT,
    "center":  WD_ALIGN_PARAGRAPH.CENTER,
    "right":   WD_ALIGN_PARAGRAPH.RIGHT,
    "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
}

PAGE_SIZES = {
    "letter": (8.5,  11.0),
    "a4":     (8.27, 11.69),
}

# ── XML helpers ───────────────────────────────────────────────

def _add_page_border(document: Document, hex_color: str):
    color    = hex_color.lstrip("#")
    sect     = document.sections[0]._sectPr
    pgBorders = OxmlElement("w:pgBorders")
    pgBorders.set(qn("w:offsetFrom"), "page")
    for side in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"),   "single")
        el.set(qn("w:sz"),    "24")
        el.set(qn("w:space"), "24")
        el.set(qn("w:color"), color)
        pgBorders.append(el)
    sect.append(pgBorders)

def _no_border_cell(cell):
    tcPr      = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in ("top","left","bottom","right","insideH","insideV"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), "none")
        tcBorders.append(el)
    tcPr.append(tcBorders)

def _bottom_border_cell(cell, color="334155", sz="6"):
    tcPr      = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    bot = OxmlElement("w:bottom")
    bot.set(qn("w:val"),   "single")
    bot.set(qn("w:sz"),    sz)
    bot.set(qn("w:color"), color)
    tcBorders.append(bot)
    tcPr.append(tcBorders)

def _shade_cell(cell, fill_hex: str):
    tcPr = cell._tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  fill_hex.lstrip("#"))
    tcPr.append(shd)

def _hex_to_rgb(hex_color: str) -> RGBColor:
    h = hex_color.lstrip("#")
    return RGBColor(int(h[0:2],16), int(h[2:4],16), int(h[4:6],16))

def _add_hr_border(paragraph, color="CBD5E1", sz="6"):
    pPr  = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot  = OxmlElement("w:bottom")
    bot.set(qn("w:val"),   "single")
    bot.set(qn("w:sz"),    sz)
    bot.set(qn("w:color"), color)
    pBdr.append(bot)
    pPr.append(pBdr)

def _add_page_number_field(paragraph):
    """Inserts a PAGE field code into paragraph."""
    run      = paragraph.add_run()
    fldBegin = OxmlElement("w:fldChar"); fldBegin.set(qn("w:fldCharType"), "begin")
    instr    = OxmlElement("w:instrText"); instr.text = " PAGE "
    fldEnd   = OxmlElement("w:fldChar"); fldEnd.set(qn("w:fldCharType"), "end")
    run._r.append(fldBegin)
    run2 = paragraph.add_run(); run2._r.append(instr)
    run3 = paragraph.add_run(); run3._r.append(fldEnd)

# ── Run builder ───────────────────────────────────────────────

def _add_runs(paragraph, runs: list, font: str, size: float):
    for rd in (runs or []):
        if rd.get("break"):
            paragraph.add_run("\n")
            continue
        if rd.get("blank"):
            r = paragraph.add_run("________________")
            r.font.name = font
            r.font.size = Pt(size)
            continue
        text = rd.get("text", "")
        r    = paragraph.add_run(text)
        r.font.name = font
        r.font.size = Pt(rd.get("size", size))
        if rd.get("bold"):      r.bold          = True
        if rd.get("italic"):    r.italic        = True
        if rd.get("underline"): r.underline     = True
        if rd.get("strike"):    r.font.strike   = True
        if rd.get("color"):     r.font.color.rgb = _hex_to_rgb(rd["color"])

# ── Block renderers ───────────────────────────────────────────

def _render_title(document: Document, json_doc: dict, font: str, size: float):
    title = json_doc.get("title")
    if not title:
        return
    title_size  = json_doc.get("title_size", round(size * 2))
    title_align = ALIGN_MAP.get(json_doc.get("title_align","center"), WD_ALIGN_PARAGRAPH.CENTER)
    for i, line in enumerate(title.split("\n")):
        p = document.add_paragraph()
        r = p.add_run(line)
        r.font.name = font
        r.font.size = Pt(title_size)
        r.bold      = True
        p.alignment = title_align
        p.paragraph_format.space_before = Pt(0) if i > 0 else Pt(4)
        p.paragraph_format.space_after  = Pt(20) if i == len(title.split("\n"))-1 else Pt(2)

def _render_heading(document: Document, block: dict, font: str, size: float):
    level = block.get("level", 1)
    p     = document.add_paragraph()
    r     = p.add_run(block.get("text",""))
    r.font.name = font
    r.font.size = Pt(size)
    r.bold      = True
    if level == 1: r.underline = True
    if level == 3: r.italic   = True
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(4)

def _render_paragraph(document: Document, block: dict, font: str, size: float):
    runs = block.get("runs", [])
    if not runs:
        return
    p = document.add_paragraph()
    _add_runs(p, runs, font, size)
    p.alignment = ALIGN_MAP.get(block.get("align","left"), WD_ALIGN_PARAGRAPH.LEFT)
    p.paragraph_format.space_before      = Pt(0)
    p.paragraph_format.space_after       = Pt(4)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    p.paragraph_format.line_spacing      = 1.35
    if block.get("indent"):
        p.paragraph_format.left_indent = Inches(block["indent"])

def _render_bullet(document: Document, block: dict, font: str, size: float):
    runs = block.get("runs", [])
    if not runs:
        return
    p      = document.add_paragraph()
    prefix = p.add_run("-  ")
    prefix.font.name = font
    prefix.font.size = Pt(size)
    _add_runs(p, runs, font, size)
    fmt = p.paragraph_format
    fmt.left_indent       = Inches(0.5)
    fmt.first_line_indent = Pt(-14)
    fmt.space_before      = Pt(0)
    fmt.space_after       = Pt(3)
    fmt.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    fmt.line_spacing      = 1.35
    p.alignment           = WD_ALIGN_PARAGRAPH.JUSTIFY

def _render_numbered(document: Document, block: dict, font: str, size: float):
    runs = block.get("runs", [])
    if not runs:
        return
    p      = document.add_paragraph()
    prefix = p.add_run(f"{block.get('index',1)}.  ")
    prefix.font.name = font
    prefix.font.size = Pt(size)
    _add_runs(p, runs, font, size)
    fmt = p.paragraph_format
    fmt.left_indent       = Inches(0.5)
    fmt.first_line_indent = Pt(-14)
    fmt.space_before      = Pt(0)
    fmt.space_after       = Pt(3)
    fmt.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    fmt.line_spacing      = 1.35

def _render_table(document: Document, block: dict, font: str, size: float):
    rows       = block.get("rows", [])
    has_header = block.get("header", True)
    if not rows:
        return
    num_cols = max(len(r) for r in rows)
    table    = document.add_table(rows=len(rows), cols=num_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit   = True
    for r_idx, row in enumerate(rows):
        is_hdr = has_header and r_idx == 0
        for c_idx in range(num_cols):
            cell      = table.rows[r_idx].cells[c_idx]
            cell.text = ""
            val       = row[c_idx] if c_idx < len(row) else ""
            run       = cell.paragraphs[0].add_run(str(val))
            run.font.name = font
            run.font.size = Pt(size)
            run.bold      = is_hdr
            cell.paragraphs[0].paragraph_format.space_before = Pt(2)
            cell.paragraphs[0].paragraph_format.space_after  = Pt(2)
            if is_hdr:
                _shade_cell(cell, "EFF6FF")

def _render_signature(document: Document, block: dict, font: str, size: float):
    labels   = block.get("labels", ["Signature"])
    num_cols = len(labels)

    sp = document.add_paragraph()
    sp.paragraph_format.space_before = Pt(24)
    sp.paragraph_format.space_after  = Pt(6)

    tbl = document.add_table(rows=4, cols=num_cols)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.autofit   = True

    for col_idx, label in enumerate(labels):
        # Row 0: signature space
        cell = tbl.rows[0].cells[col_idx]
        _no_border_cell(cell)
        _bottom_border_cell(cell)
        cell.paragraphs[0].paragraph_format.space_before = Pt(40)
        cell.paragraphs[0].paragraph_format.space_after  = Pt(4)

        # Row 1: label
        cell = tbl.rows[1].cells[col_idx]
        _no_border_cell(cell)
        r = cell.paragraphs[0].add_run(label)
        r.bold = True; r.font.name = font; r.font.size = Pt(size - 1)
        cell.paragraphs[0].paragraph_format.space_before = Pt(2)
        cell.paragraphs[0].paragraph_format.space_after  = Pt(10)

        # Row 2: name/date blank line
        cell = tbl.rows[2].cells[col_idx]
        _no_border_cell(cell)
        _bottom_border_cell(cell)
        cell.paragraphs[0].paragraph_format.space_before = Pt(28)
        cell.paragraphs[0].paragraph_format.space_after  = Pt(4)

        # Row 3: "Name / Date" label
        cell = tbl.rows[3].cells[col_idx]
        _no_border_cell(cell)
        r = cell.paragraphs[0].add_run("Name / Date")
        r.font.name      = font
        r.font.size      = Pt(size - 2)
        r.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)
        cell.paragraphs[0].paragraph_format.space_before = Pt(2)

def _render_hr(document: Document):
    p = document.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(8)
    _add_hr_border(p)

def _render_spacer(document: Document, block: dict):
    p = document.add_paragraph()
    p.paragraph_format.space_before = Pt(block.get("size", 12))
    p.paragraph_format.space_after  = Pt(0)

def _render_pagebreak(document: Document):
    p   = document.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pb  = OxmlElement("w:pageBreakBefore")
    pb.set(qn("w:val"), "true")
    pPr.append(pb)

def _render_header_footer(document: Document, header_text: Optional[str], footer_text: Optional[str], font: str, size: float):
    section = document.sections[0]
    if header_text:
        h  = section.header
        p  = h.paragraphs[0] if h.paragraphs else h.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r  = p.add_run(header_text)
        r.font.name      = font
        r.font.size      = Pt(size - 2)
        r.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)
        _add_hr_border(p, color="CBD5E1")

    if footer_text:
        f  = section.footer
        p  = f.paragraphs[0] if f.paragraphs else f.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        parts = re.split(r'(\{PAGE\})', footer_text)
        for part in parts:
            if part == "{PAGE}":
                _add_page_number_field(p)
            elif part:
                r = p.add_run(part)
                r.font.name      = font
                r.font.size      = Pt(size - 2)
                r.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)
        _add_hr_border(p, color="CBD5E1")


# ─────────────────────────────────────────────────────────────
# MAIN RENDER FUNCTION  — takes engine JSON → bytes
# ─────────────────────────────────────────────────────────────

def render_document(json_doc: dict) -> bytes:
    """
    Universal deterministic renderer.
    Input:  engine JSON (from stage3 or /generate-docx)
    Output: .docx bytes
    ZERO hallucination — only renders what's in the JSON.
    Unknown block types are silently skipped.
    """
    font     = json_doc.get("font",      "Times New Roman")
    size     = float(json_doc.get("font_size", 12))
    page_key = json_doc.get("page_size", "letter").lower()
    w, h     = PAGE_SIZES.get(page_key, PAGE_SIZES["letter"])
    margins  = json_doc.get("margins", [1, 1, 1, 1])

    document = Document()

    # Page setup
    section              = document.sections[0]
    section.page_width   = Inches(w)
    section.page_height  = Inches(h)
    section.top_margin   = Inches(margins[0])
    section.right_margin = Inches(margins[1])
    section.bottom_margin= Inches(margins[2])
    section.left_margin  = Inches(margins[3])

    # Default style
    try:
        document.styles["Normal"].font.name = font
        document.styles["Normal"].font.size = Pt(size)
    except Exception:
        pass

    # Page border
    if json_doc.get("page_border"):
        _add_page_border(document, json_doc["page_border"])

    # Header / footer
    _render_header_footer(
        document,
        json_doc.get("header_text"),
        json_doc.get("footer_text"),
        font, size
    )

    # Title
    _render_title(document, json_doc, font, size)

    # Content blocks
    for block in json_doc.get("content", []):
        btype = block.get("type")

        if   btype == "heading":   _render_heading(document, block, font, size)
        elif btype == "paragraph": _render_paragraph(document, block, font, size)
        elif btype == "bullet":    _render_bullet(document, block, font, size)
        elif btype == "numbered":  _render_numbered(document, block, font, size)
        elif btype == "table":     _render_table(document, block, font, size)
        elif btype == "signature": _render_signature(document, block, font, size)
        elif btype == "hr":        _render_hr(document)
        elif btype == "spacer":    _render_spacer(document, block)
        elif btype == "pagebreak": _render_pagebreak(document)
        # Unknown → silently skip. NEVER invent content.

    buf = io.BytesIO()
    document.save(buf)
    buf.seek(0)
    return buf.read()


# =============================================================
# ░░░░  SECTION 3 — BRAND HEADER (for offer/NOC layouts) ░░░░░
# =============================================================

def fetch_image(url: Optional[str]) -> Optional[io.BytesIO]:
    if not url:
        return None
    try:
        if url.startswith("http"):
            r = requests.get(url, timeout=8); r.raise_for_status()
            return io.BytesIO(r.content)
        local = os.path.join(BASE_DIR, "public", url.lstrip("/"))
        if os.path.exists(local):
            with open(local, "rb") as f:
                return io.BytesIO(f.read())
    except Exception as e:
        print(f"⚠️ fetch_image failed: {e}")
    return None

DOC_LAYOUTS: dict = {
    "default":              {"showLogo": False, "showSignature": False, "headerImageUrl": None, "footerImageUrl": None},
    "offer_modern_blue":    {"showLogo": True,  "showSignature": True,  "headerImageUrl": "/graphics/offer/header-mod-blue.png",   "footerImageUrl": "/graphics/offer/footer-wave-blue.png"},
    "offer_green_wave":     {"showLogo": True,  "showSignature": True,  "headerImageUrl": "/graphics/offer/header-green-wave.webp", "footerImageUrl": "/graphics/offer/footer-green-wave.webp"},
    "offer_minimal_plain":  {"showLogo": True,  "showSignature": True,  "headerImageUrl": None, "footerImageUrl": None},
    "offer_classic_border": {"showLogo": True,  "showSignature": True,  "headerImageUrl": None, "footerImageUrl": None},
    "noc_plain":            {"showLogo": True,  "showSignature": True,  "headerImageUrl": None, "footerImageUrl": None},
    "rental_plain":         {"showLogo": True,  "showSignature": True,  "headerImageUrl": None, "footerImageUrl": None},
    "plain_editor":         {"showLogo": False, "showSignature": False, "headerImageUrl": None, "footerImageUrl": None},
}

SLUG_OVERRIDES: dict = {
    "visa-expiration-letter":       "plain_editor",
    "website-proposal-standard":    "plain_editor",
    "mobile-app-proposal-standard": "plain_editor",
    "blog-article-standard":        "plain_editor",
    "rental-agreement-11-months":   "rental_plain",
    "offer-letter-standard":        "offer_modern_blue",
    "noc-employee-visa":            "noc_plain",
}

def get_layout(template_slug: Optional[str], design_key: Optional[str]) -> dict:
    s = (template_slug or "").lower()
    if s.startswith("leave-application-"):  return DOC_LAYOUTS["plain_editor"]
    if design_key and design_key in DOC_LAYOUTS: return DOC_LAYOUTS[design_key]
    if template_slug:
        if template_slug in SLUG_OVERRIDES:      return DOC_LAYOUTS[SLUG_OVERRIDES[template_slug]]
        if any(k in s for k in ("offer","appointment","joining")): return DOC_LAYOUTS["offer_modern_blue"]
        if "noc" in s:                           return DOC_LAYOUTS["noc_plain"]
        if any(k in s for k in ("rental","lease")): return DOC_LAYOUTS["rental_plain"]
        if any(k in s for k in ("blog","ai-blog","content","copywriter","proposal")): return DOC_LAYOUTS["plain_editor"]
    return DOC_LAYOUTS["default"]

def render_brand_header(document: Document, layout: dict, brand: Optional[dict], title: Optional[str]):
    BODY_FONT = "Times New Roman"
    header_img = fetch_image(layout.get("headerImageUrl"))
    if header_img:
        p = document.add_paragraph()
        p.add_run().add_picture(header_img, width=Inches(6.5))
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(6)

    if not brand:
        return

    from docx.oxml import OxmlElement as Oxm
    tbl = document.add_table(rows=1, cols=3)
    tbl.autofit = True
    lc, cc, rc = tbl.rows[0].cells
    for cell in (lc, cc, rc):
        tcPr = cell._tc.get_or_add_tcPr()
        tcB  = Oxm("w:tcBorders")
        for s in ("top","left","bottom","right","insideH","insideV"):
            e = Oxm(f"w:{s}"); e.set(qn("w:val"),"none"); tcB.append(e)
        tcPr.append(tcB)

    logo_img = fetch_image(brand.get("logoUrl")) if layout.get("showLogo") else None
    if logo_img:
        lc.paragraphs[0].add_run().add_picture(logo_img, width=Inches(1.2))

    nr = cc.paragraphs[0].add_run(brand.get("companyName",""))
    nr.bold = True; nr.font.size = Pt(10); nr.font.name = BODY_FONT
    for line in (brand.get("addressLine1"), brand.get("addressLine2")):
        if line:
            p = cc.add_paragraph(line)
            for r in p.runs: r.font.size = Pt(8); r.font.name = BODY_FONT; r.font.color.rgb = RGBColor(0x47,0x55,0x69)

    for val in (brand.get("phone"), brand.get("email")):
        if val:
            p = rc.add_paragraph(val)
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            for r in p.runs: r.font.size = Pt(8); r.font.name = BODY_FONT; r.font.color.rgb = RGBColor(0x47,0x55,0x69)

    if title:
        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(title.upper())
        run.bold = True; run.font.size = Pt(9); run.font.name = BODY_FONT
        run.font.color.rgb = RGBColor(0x1D,0x4E,0xD8)
        pPr = p._p.get_or_add_pPr(); pBdr = OxmlElement("w:pBdr")
        for side in ("top","bottom"):
            b = OxmlElement(f"w:{side}")
            b.set(qn("w:val"),"single"); b.set(qn("w:sz"),"8"); b.set(qn("w:color"),"3B82F6")
            pBdr.append(b)
        pPr.append(pBdr)
        p.paragraph_format.space_before = Pt(6); p.paragraph_format.space_after = Pt(8)
    document.add_paragraph()

def render_signatory_footer(document: Document, signatory: dict):
    BODY_FONT = "Times New Roman"
    sp = document.add_paragraph()
    sp.paragraph_format.space_before = Pt(24); sp.paragraph_format.space_after = Pt(4)
    lp = document.add_paragraph("Authorised Signatory")
    for r in lp.runs: r.bold=True; r.font.size=Pt(9); r.font.name=BODY_FONT; r.font.color.rgb=RGBColor(0x47,0x55,0x69)
    sig_img = fetch_image(signatory.get("signatureImageUrl"))
    if sig_img:
        p = document.add_paragraph(); p.add_run().add_picture(sig_img, width=Inches(1.5))
    else:
        blank = document.add_paragraph(); blank.paragraph_format.space_before = Pt(16)
    for text, bold in ((signatory.get("fullName",""),True),(signatory.get("designation",""),False)):
        if text:
            p = document.add_paragraph(text)
            for r in p.runs:
                r.bold=bold; r.font.size=Pt(12 if bold else 10); r.font.name=BODY_FONT
                if not bold: r.font.color.rgb=RGBColor(0x47,0x55,0x69)

def render_footer_banner(document: Document, layout: dict):
    footer_img = fetch_image(layout.get("footerImageUrl"))
    if not footer_img: return
    p = document.add_paragraph()
    p.add_run().add_picture(footer_img, width=Inches(6.5))
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)


def sanitize_filename(name: str) -> str:
    name = unicodedata.normalize("NFKD", name).encode("ascii","ignore").decode("ascii")
    name = re.sub(r"[^\w\s\-.]", "_", name)
    return name.strip() or "document"


# =============================================================
# ░░░░  SECTION 4 — ALL ROUTES  ░░░░░░░░░░░░░░░░░░░░░░░░░░░░░
# =============================================================

class ProcessRequest(BaseModel):
    jobId: str

@app.post("/api/handwritten/process")
async def start_handwritten_process(payload: ProcessRequest, background_tasks: BackgroundTasks):
    log("Starting OCR job", payload.jobId)
    update_job(payload.jobId, state="queued")
    background_tasks.add_task(run_ocr_job, payload.jobId)
    return {"started": True}

@app.post("/api/job-register")
async def register_job(payload: dict):
    jobId = payload["jobId"]
    update_job(jobId, filePath=payload["filePath"],
               source=payload.get("source","scanned"),
               strict=payload.get("strict", True), state="uploaded")
    return {"ok": True}

@app.get("/api/job-status")
async def job_status(jobId: str):
    job = load_job(jobId)
    if not job:
        raise HTTPException(status_code=404, detail="Job not found")
    return job

@app.post("/api/job-complete-free")
async def complete_free_job(payload: dict):
    update_job(payload["jobId"], state="free-ready", source="digital-pdf")
    return {"ok": True}

@app.post("/api/detect-pdf-type")
async def detect_pdf_type_route(file: UploadFile = File(...)):
    data = await file.read()
    doc  = fitz.open(stream=data, filetype="pdf")
    for i in range(min(len(doc), 3)):
        if doc.load_page(i).get_text().strip():
            return {"type": "digital"}
    return {"type": "scanned"}

class ExportRequest(BaseModel):
    filePath: str

@app.post("/api/export-digital-docx")
async def export_digital_docx(payload: ExportRequest):
    if not os.path.exists(payload.filePath):
        raise HTTPException(status_code=400, detail="FILE_NOT_FOUND")
    with open(payload.filePath,"rb") as f:
        pdf_bytes = f.read()
    pdf      = fitz.open(stream=pdf_bytes, filetype="pdf")
    word_doc = Document()
    s = word_doc.sections[0]
    s.top_margin = s.bottom_margin = s.left_margin = s.right_margin = Inches(1)
    for page in pdf:
        raw_text = page.get_text().strip()
        if not raw_text: continue
        for block in [b.strip() for b in raw_text.split("\n\n") if b.strip()]:
            p = word_doc.add_paragraph(block)
            p.paragraph_format.line_spacing = 1.5
            p.paragraph_format.space_after  = Pt(12)
            for run in p.runs: run.font.name="Times New Roman"; run.font.size=Pt(12)
    buf = io.BytesIO(); word_doc.save(buf); buf.seek(0)
    return StreamingResponse(buf,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": "attachment; filename=Converted_Document.docx"})

@app.post("/api/upload")
async def upload_file(file: UploadFile = File(...)):
    try:
        os.makedirs("uploads", exist_ok=True)
        file_path = os.path.join("uploads", f"{datetime.now().timestamp()}_{file.filename}")
        with open(file_path,"wb") as f:
            f.write(await file.read())
        return {"filePath": file_path}
    except Exception:
        traceback.print_exc()
        raise HTTPException(status_code=500, detail="UPLOAD_FAILED")

@app.post("/api/parse-document")
async def parse_document_route(
    file:   UploadFile = File(...),
    strict: bool       = Form(True),
    source: str        = Form("scanned"),
):
    log("API HIT /api/parse-document")
    try:
        raw_bytes = await file.read()
        if not raw_bytes:
            raise HTTPException(status_code=400, detail="EMPTY_FILE")
        image_bytes = pdf_to_image_bytes(raw_bytes) if is_pdf(raw_bytes) else to_png_bytes(raw_bytes)
        document    = parse_document(image_bytes)
        return {"success": True, "engine_version": ENGINE_VERSION, "document": document}
    except HTTPException:
        raise
    except Exception as e:
        log("❌ ROUTE ERROR", repr(e)); traceback.print_exc()
        raise HTTPException(status_code=500, detail="PROCESSING_FAILED")


# ── /generate-docx — Universal renderer entry point ──────────

class GenerateDocxRequest(BaseModel):
    contentJson:  Any           = None
    fileName:     Optional[str] = Field(default="document")
    templateSlug: Optional[str] = None
    designKey:    Optional[str] = None
    brand:        Optional[Any] = None
    signatory:    Optional[Any] = None
    baseTemplate: Optional[str] = None

@app.post("/generate-docx")
async def generate_docx_route(payload: GenerateDocxRequest):
    try:
        log("GENERATE DOCX", f"slug={payload.templateSlug} design={payload.designKey}")

        content = payload.contentJson or {}

        # ── Path A: NEW engine JSON (has "content" array with typed blocks)
        if isinstance(content, dict) and isinstance(content.get("content"), list):
            docx_bytes = render_document(content)

            # If layout needs brand header/footer, wrap with it
            layout = get_layout(payload.templateSlug, payload.designKey)
            if layout.get("showLogo") or layout.get("headerImageUrl") or layout.get("footerImageUrl"):
                # Rebuild with brand header injected
                inner_doc = Document()
                section   = inner_doc.sections[0]
                section.top_margin = section.bottom_margin = section.left_margin = section.right_margin = Inches(1)
                if layout.get("showLogo") or layout.get("headerImageUrl"):
                    render_brand_header(inner_doc, layout, payload.brand, payload.fileName)
                # Re-render content blocks into inner_doc
                font = content.get("font","Times New Roman")
                size = float(content.get("font_size",12))
                for block in content.get("content",[]):
                    btype = block.get("type")
                    if   btype == "heading":   _render_heading(inner_doc, block, font, size)
                    elif btype == "paragraph": _render_paragraph(inner_doc, block, font, size)
                    elif btype == "bullet":    _render_bullet(inner_doc, block, font, size)
                    elif btype == "numbered":  _render_numbered(inner_doc, block, font, size)
                    elif btype == "table":     _render_table(inner_doc, block, font, size)
                    elif btype == "signature": _render_signature(inner_doc, block, font, size)
                    elif btype == "hr":        _render_hr(inner_doc)
                    elif btype == "spacer":    _render_spacer(inner_doc, block)
                    elif btype == "pagebreak": _render_pagebreak(inner_doc)
                if layout.get("showSignature") and payload.signatory:
                    render_signatory_footer(inner_doc, payload.signatory)
                if layout.get("footerImageUrl"):
                    render_footer_banner(inner_doc, layout)
                buf = io.BytesIO(); inner_doc.save(buf); buf.seek(0)
                docx_bytes = buf.read()

        # ── Path B: Legacy TipTap JSON (backwards compatibility)
        else:
            from docx import Document as LegacyDoc
            document = LegacyDoc()
            section  = document.sections[0]
            section.top_margin = section.bottom_margin = section.left_margin = section.right_margin = Inches(1)
            layout = get_layout(payload.templateSlug, payload.designKey)
            if layout.get("showLogo") or layout.get("headerImageUrl"):
                render_brand_header(document, layout, payload.brand, payload.fileName)
            # minimal TipTap fallback — render paragraphs only
            if isinstance(content, dict) and content.get("type") == "doc":
                for node in content.get("content",[]):
                    ntype = node.get("type")
                    if ntype in ("paragraph","heading"):
                        texts = [c.get("text","") for c in node.get("content",[]) if c.get("type")=="text"]
                        if texts:
                            p = document.add_paragraph(" ".join(texts))
                            p.paragraph_format.space_after = Pt(6)
            if layout.get("showSignature") and payload.signatory:
                render_signatory_footer(document, payload.signatory)
            buf = io.BytesIO(); document.save(buf); buf.seek(0)
            docx_bytes = buf.read()

        safe_name = sanitize_filename(payload.fileName or "document")
        if not safe_name.lower().endswith(".docx"):
            safe_name += ".docx"

        return Response(
            content=docx_bytes,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f'attachment; filename="{safe_name}"'},
        )

    except Exception as e:
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=str(e))


if __name__ == "__main__":
    import uvicorn
    uvicorn.run("app:app", host="0.0.0.0", port=8000, reload=True)